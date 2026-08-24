from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "cv"
OUT.mkdir(exist_ok=True)

DOCX_PATH = OUT / "Khanh_Huyen_UIUX_CV.docx"

NAVY = "101828"
BLUE = "2563EB"
CYAN = "0EA5E9"
MINT = "DDFBF0"
LIGHT = "F7FAFC"
GRAY = "667085"
PALE_BLUE = "EFF6FF"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="E4E7EC", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def set_font(run, size=10, color=NAVY, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def clear_para(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.05


def add_text(cell, text, size=9.5, color=NAVY, bold=False, italic=False, after=4, align=None):
    p = cell.add_paragraph()
    clear_para(p)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(cell, text, color=BLUE):
    p = cell.add_paragraph()
    clear_para(p)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    set_font(run, size=9, color=color, bold=True)
    return p


def add_bullet(cell, text, size=8.8, color=NAVY):
    p = cell.add_paragraph(style=None)
    clear_para(p)
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_after = Pt(3)
    marker = p.add_run("- ")
    set_font(marker, size=size, color=BLUE, bold=True)
    run = p.add_run(text)
    set_font(run, size=size, color=color)


def add_project(cell, title, details):
    p = cell.add_paragraph()
    clear_para(p)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    set_font(run, size=10.2, color=NAVY, bold=True)
    for detail in details:
        add_bullet(cell, detail, size=8.6)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    styles["Normal"].font.size = Pt(9.4)

    shell = doc.add_table(rows=1, cols=2)
    set_table_width(shell, [Inches(2.05), Inches(5.55)])
    shell.rows[0].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    shell.rows[0].cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    left, right = shell.rows[0].cells
    for cell in (left, right):
        set_cell_borders(cell, color="FFFFFF", size="0")
        set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
    set_cell_shading(left, PALE_BLUE)
    set_cell_shading(right, WHITE)

    # Sidebar
    add_text(left, "KHANH\nHUYEN", size=21, color=NAVY, bold=True, after=0)
    add_text(left, "UI/UX DESIGNER INTERN", size=8.5, color=BLUE, bold=True, after=10)
    add_text(left, "Ho Chi Minh City, Vietnam", size=8.7, color=GRAY)
    add_text(left, "tkhanhhuyen0484@gmail.com", size=8.2, color=NAVY)
    add_text(left, "Portfolio: add public link", size=8.2, color=NAVY)
    add_text(left, "Phone: add phone number", size=8.2, color=NAVY, after=10)

    add_heading(left, "Core Skills", color=BLUE)
    for skill in [
        "Mobile App UI Design",
        "User Flow & Wireframing",
        "High-fidelity Prototyping",
        "Visual Hierarchy",
        "Design Systems",
        "Responsive Web Design",
        "Usability Testing",
        "Developer Handoff",
    ]:
        add_bullet(left, skill, size=8.3, color=NAVY)

    add_heading(left, "Tools", color=BLUE)
    for tool in ["Figma", "Adobe XD", "Photoshop", "Illustrator", "Framer", "Webflow"]:
        add_bullet(left, tool, size=8.3, color=NAVY)

    add_heading(left, "Development", color=BLUE)
    for dev in ["HTML/CSS", "JavaScript", "React", "Tailwind CSS", "Git"]:
        add_bullet(left, dev, size=8.3, color=NAVY)

    add_heading(left, "Education", color=BLUE)
    add_text(left, "Industrial University of Ho Chi Minh City", size=8.7, color=NAVY, bold=True, after=2)
    add_text(left, "Information Technology - 4th-year student", size=8.3, color=GRAY, after=0)

    # Main column
    title = right.add_paragraph()
    clear_para(title)
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Khanh Huyen")
    set_font(run, size=25, color=NAVY, bold=True)

    sub = right.add_paragraph()
    clear_para(sub)
    sub.paragraph_format.space_after = Pt(10)
    run = sub.add_run("UI/UX Designer focused on polished mobile experiences")
    set_font(run, size=10.8, color=BLUE, bold=True)

    summary = (
        "4th-year IT student and UI/UX Design Intern based in Ho Chi Minh City. "
        "I design mobile apps, websites and product flows that feel clear, polished and easy to use. "
        "My portfolio focuses on practical app experiences across camera, AI notes, parental control, e-commerce and entertainment products."
    )
    add_text(right, summary, size=9.5, color=NAVY, after=8)

    add_heading(right, "Selected Projects", color=BLUE)
    add_project(
        right,
        "KidGuard App - Parental Control Mobile App",
        [
            "Designed parent and child flows for screen time, app access, setup and lock states.",
            "Created a friendly visual system using soft blue, green and safety-focused UI patterns.",
        ],
    )
    add_project(
        right,
        "AI Voice Lock: Smart Note - Secure Notes App",
        [
            "Designed voice lock, AI note list, onboarding, setup and private preview flows.",
            "Built a security-focused interface with clear permissions, lock modes and summary actions.",
        ],
    )
    add_project(
        right,
        "Ultra Camera - Mobile Camera App",
        [
            "Designed capture, gallery, onboarding and photo editing screens for a pro camera flow.",
            "Applied a clean purple UI system with clear primary actions and editing controls.",
        ],
    )
    add_project(
        right,
        "ELOGI - Mobile E-commerce App",
        [
            "Designed product discovery, search, product detail, chat, notification and delivery tracking screens.",
            "Created a clean mobile shopping experience with strong CTA hierarchy and consistent components.",
        ],
    )
    add_project(
        right,
        "Prank Sound - Mobile Entertainment App",
        [
            "Designed a playful app experience for browsing and playing prank sound effects.",
            "Organized sound categories, player actions and rating moments into a bold, easy-to-tap UI.",
        ],
    )

    add_heading(right, "Experience & Strengths", color=BLUE)
    for item in [
        "Create user-centered mobile screens from research, user flow and wireframe through high-fidelity UI.",
        "Use spacing, typography, color and component consistency to make interfaces feel calm and premium.",
        "Prepare organized screens, states and reusable UI patterns for smoother developer handoff.",
    ]:
        add_bullet(right, item, size=8.8)

    add_heading(right, "Profile Keywords", color=BLUE)
    add_text(
        right,
        "UI/UX Design - Product Design - Mobile App Design - Figma - Prototyping - User Flow - Visual Design - Design Systems",
        size=8.8,
        color=GRAY,
        after=0,
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_para(footer)
    fr = footer.add_run("Khanh Huyen - UI/UX Designer Intern")
    set_font(fr, size=8, color=GRAY)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build())
