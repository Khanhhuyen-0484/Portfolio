from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "cv"
PDF_PATH = OUT / "Khanh_Huyen_CV_Hoan_Chinh.pdf"
DOCX_PATH = OUT / "Khanh_Huyen_CV_Hoan_Chinh.docx"
PORTRAIT = ROOT / "img" / "about-portrait.png"

pdfmetrics.registerFont(TTFont("ArialVN", "C:/Windows/Fonts/arial.ttf"))
pdfmetrics.registerFont(TTFont("ArialVNBold", "C:/Windows/Fonts/arialbd.ttf"))

INK = colors.HexColor("#26262E")
TEXT = colors.HexColor("#2F333A")
MUTED = colors.HexColor("#666B73")
SOFT = colors.HexColor("#E7E8E6")
PANEL = colors.HexColor("#E3E4E1")
RULE = colors.HexColor("#9B9D9F")
ACCENT = colors.HexColor("#212129")


PROFILE = {
    "name": "KHÁNH HUYỀN",
    "role": "UI/UX DESIGNER",
    "phone": "0328895929",
    "email": "tkhanhhuyen0484@gmail.com",
    "location": "Gò Vấp, TP. Hồ Chí Minh",
    "website": "huyensaigon.id.vn",
}

ABOUT = (
    "Sinh viên năm 4 ngành Công nghệ thông tin, định hướng UI/UX Design. "
    "Có kinh nghiệm thực tập tại EMC và làm freelance, tập trung thiết kế app, "
    "website, landing page, banner và prototype."
)

EDUCATION = [
    ("Đại học Công nghiệp TP. Hồ Chí Minh", "Công nghệ thông tin", "2022 - nay"),
]

SKILLS = [
    ("UI mobile/web", 0.90),
    ("Wireframe", 0.84),
    ("Prototype", 0.88),
    ("Visual design", 0.82),
    ("Banner / LP", 0.78),
    ("Handoff", 0.74),
]

TOOLS = ["Figma", "Adobe XD", "Photoshop", "Illustrator", "HTML/CSS", "JavaScript"]

LANGUAGES = ["Tiếng Việt", "Tiếng Anh đọc hiểu tài liệu chuyên ngành"]

EXPERIENCE = [
    (
        "UI/UX Designer Intern",
        "EMC",
        "06/2025 - 08/2025",
        [
            "Thiết kế giao diện app, website và landing page theo yêu cầu dự án.",
            "Chỉnh UI, tối ưu bố cục, màu sắc, typography và trạng thái giao diện.",
            "Thiết kế banner, visual assets và prototype tương tác phục vụ demo sản phẩm.",
        ],
    ),
    (
        "Freelance UI/UX Designer",
        "Freelance",
        "08/2025 - nay",
        [
            "Nhận thiết kế giao diện mobile app, website, landing page và các ấn phẩm số.",
            "Tư vấn bố cục, visual direction và cải thiện UI để sản phẩm chuyên nghiệp hơn.",
            "Chuẩn bị prototype, màn hình thiết kế và tài liệu bàn giao theo nhu cầu khách hàng.",
        ],
    ),
]

PRACTICE = [
    (
        "Thiết kế sản phẩm số",
        "Thực hành thiết kế các luồng onboarding, home, setting, gallery, edit flow và lock state cho mobile app.",
    ),
    (
        "UI system cơ bản",
        "Sắp xếp component, màu sắc, typography và state UI có hệ thống để hỗ trợ quá trình handoff.",
    ),
]


def pdf_font(bold=False):
    return "ArialVNBold" if bold else "ArialVN"


def wrap_text(value, font_name, size, width):
    words = value.split()
    lines, line = [], ""
    for word in words:
        trial = word if not line else f"{line} {word}"
        if stringWidth(trial, font_name, size) <= width:
            line = trial
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def draw_text(c, x, y, value, size=8.5, color=TEXT, bold=False, leading=10.5, width=None):
    font_name = pdf_font(bold)
    c.setFont(font_name, size)
    c.setFillColor(color)
    lines = wrap_text(value, font_name, size, width) if width else value.split("\n")
    for line in lines:
        c.drawString(x, y, line)
        y -= leading
    return y


def section_title(c, x, y, title, width):
    c.setFillColor(ACCENT)
    c.setFont(pdf_font(True), 8.3)
    label = title.upper()
    c.drawString(x, y, label)
    c.setStrokeColor(RULE)
    c.setLineWidth(0.7)
    line_start = x + stringWidth(label, pdf_font(True), 8.3) + 9
    c.line(line_start, y + 3, x + width, y + 3)
    return y - 17


def tiny_icon(c, x, y):
    c.setStrokeColor(ACCENT)
    c.setLineWidth(0.8)
    c.rect(x, y - 5, 7, 7, stroke=1, fill=0)
    c.setFillColor(ACCENT)
    c.circle(x + 3.5, y - 1.5, 1.1, stroke=0, fill=1)


def bullet(c, x, y, text, width, size=7.3, leading=9.0):
    c.setStrokeColor(ACCENT)
    c.setLineWidth(0.8)
    c.circle(x + 2.5, y + 2.5, 2.2, stroke=1, fill=0)
    return draw_text(c, x + 12, y, text, size=size, color=TEXT, leading=leading, width=width - 12) - 2


def draw_portrait(c, x, y, size):
    if not PORTRAIT.exists():
        c.setFillColor(colors.white)
        c.circle(x + size / 2, y + size / 2, size / 2, fill=1, stroke=0)
        return
    c.saveState()
    path = c.beginPath()
    path.circle(x + size / 2, y + size / 2, size / 2)
    c.clipPath(path, stroke=0, fill=0)
    c.drawImage(str(PORTRAIT), x, y, width=size, height=size, preserveAspectRatio=True, anchor="c")
    c.restoreState()
    c.setStrokeColor(colors.white)
    c.setLineWidth(2.5)
    c.circle(x + size / 2, y + size / 2, size / 2, stroke=1, fill=0)


def draw_skill_bar(c, x, y, label, level, width):
    draw_text(c, x, y, label, 7.3, TEXT, False, 9, width * 0.58)
    bar_x = x + width * 0.62
    bar_y = y + 2
    bar_w = width * 0.30
    c.setStrokeColor(colors.HexColor("#C1C3C5"))
    c.setLineWidth(2.0)
    c.line(bar_x, bar_y, bar_x + bar_w, bar_y)
    c.setStrokeColor(ACCENT)
    c.line(bar_x, bar_y, bar_x + bar_w * level, bar_y)
    return y - 13


def draw_pdf():
    OUT.mkdir(exist_ok=True)
    W, H = letter
    page_x = 0.72 * inch
    page_y = 0.50 * inch
    page_w = W - 2 * page_x
    page_h = H - 0.94 * inch
    side_w = 1.92 * inch
    main_x = page_x + side_w
    main_w = page_w - side_w

    c = canvas.Canvas(str(PDF_PATH), pagesize=letter)
    c.setTitle("Khánh Huyền CV")

    c.setFillColor(colors.white)
    c.rect(0, 0, W, H, fill=1, stroke=0)
    c.setFillColor(PANEL)
    c.rect(page_x, page_y, side_w, page_h, fill=1, stroke=0)

    portrait_size = 1.34 * inch
    draw_portrait(c, page_x + 0.32 * inch, H - 1.82 * inch, portrait_size)

    header_y = H - 1.46 * inch
    header_h = 0.88 * inch
    c.setFillColor(INK)
    c.rect(main_x - 0.02 * inch, header_y, main_w + 0.02 * inch, header_h, fill=1, stroke=0)
    draw_text(c, main_x + 0.24 * inch, header_y + 0.51 * inch, PROFILE["name"], 17.5, colors.white, True, 19)
    c.setFont(pdf_font(True), 7.2)
    c.setFillColor(colors.white)
    c.drawString(main_x + 0.25 * inch, header_y + 0.25 * inch, PROFILE["role"])

    sx = page_x + 0.20 * inch
    sy = H - 2.60 * inch
    sw = side_w - 0.42 * inch
    sy = section_title(c, sx, sy, "About me", sw)
    sy = draw_text(c, sx, sy, ABOUT, 6.65, TEXT, False, 8.0, sw)
    sy -= 16

    sy = section_title(c, sx, sy, "Education", sw)
    for school, major, period in EDUCATION:
        sy = draw_text(c, sx, sy, school, 6.9, TEXT, True, 8.3, sw)
        sy = draw_text(c, sx, sy + 1, major, 6.55, MUTED, False, 8.0, sw)
        sy = draw_text(c, sx, sy + 1, period, 6.55, MUTED, False, 8.0, sw)
        sy -= 10

    sy = section_title(c, sx, sy, "Skills", sw)
    for label, level in SKILLS:
        sy = draw_skill_bar(c, sx, sy, label, level, sw)
    sy -= 7

    sy = section_title(c, sx, sy, "Tools", sw)
    for tool in TOOLS:
        sy = bullet(c, sx, sy, tool, sw, 6.75, 8.1)
    sy -= 7

    sy = section_title(c, sx, sy, "Language", sw)
    for lang in LANGUAGES:
        sy = bullet(c, sx, sy, lang, sw, 6.75, 8.1)

    mx = main_x + 0.23 * inch
    my = header_y - 0.42 * inch
    mw = main_w - 0.44 * inch
    contact_gap = mw / 2
    contacts = [
        (PROFILE["phone"], PROFILE["website"]),
        (PROFILE["email"], PROFILE["location"]),
    ]
    for row, pair in enumerate(contacts):
        y = my - row * 0.24 * inch
        for col, value in enumerate(pair):
            x = mx + col * contact_gap
            tiny_icon(c, x, y + 3)
            draw_text(c, x + 0.14 * inch, y, value, 6.35, MUTED, False, 7.2, contact_gap - 0.2 * inch)

    my -= 0.72 * inch
    my = section_title(c, mx, my, "Experience", mw)
    timeline_x = mx + 0.04 * inch
    c.setStrokeColor(RULE)
    c.setLineWidth(0.7)
    c.line(timeline_x, my + 11, timeline_x, my - 2.15 * inch)
    for title, org, period, items in EXPERIENCE:
        c.setStrokeColor(ACCENT)
        c.circle(timeline_x, my + 2, 2.6, stroke=1, fill=0)
        draw_text(c, mx + 0.20 * inch, my, title, 8.2, TEXT, True, 9.7, mw * 0.58)
        draw_text(c, mx + mw - 1.02 * inch, my, period, 6.65, MUTED, False, 8, 1.02 * inch)
        my = draw_text(c, mx + 0.20 * inch, my - 10, org, 6.9, MUTED, False, 8, mw - 0.2 * inch)
        for item in items:
            my = bullet(c, mx + 0.27 * inch, my + 1, item, mw - 0.36 * inch, 6.85, 8.35)
        my -= 9

    my = section_title(c, mx, my + 1, "Practice", mw)
    for title, desc in PRACTICE:
        draw_text(c, mx, my, title, 7.8, TEXT, True, 9.2, mw)
        my = draw_text(c, mx, my - 10, desc, 6.9, MUTED, False, 8.4, mw)
        my -= 8

    my = section_title(c, mx, my + 1, "Strengths", mw)
    for item in [
        "Quan sát chi tiết, yêu thích sự gọn gàng và nhất quán trong giao diện.",
        "Chủ động học công cụ mới, tiếp thu phản hồi và cải thiện thiết kế qua từng vòng chỉnh sửa.",
        "Có nền tảng kỹ thuật giúp giao tiếp tốt hơn với developer trong quá trình làm sản phẩm.",
    ]:
        my = bullet(c, mx, my, item, mw, 6.9, 8.5)

    c.save()


def set_cell_bg(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_border(cell, color="FFFFFF", size="0"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def run_style(run, size=9, color="2F333A", bold=False, uppercase=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.text = run.text.upper() if uppercase else run.text


def add_docx_text(cell, value, size=8.5, color="2F333A", bold=False, after=3, align=None):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    if align:
        p.alignment = align
    r = p.add_run(value)
    run_style(r, size, color, bold)
    return p


def docx_heading(cell, title):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title.upper())
    run_style(r, 8.5, "26262E", True)


def docx_bullet(cell, value, size=8.2):
    p = cell.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.first_line_indent = Inches(-0.10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("- " + value)
    run_style(r, size, "2F333A")


def clear_paragraph(cell):
    first = cell.paragraphs[0]
    first._element.getparent().remove(first._element)


def draw_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Inches(2.05)
    right.width = Inches(5.0)
    set_cell_bg(left, "E3E4E1")
    set_cell_bg(right, "FFFFFF")
    set_cell_border(left)
    set_cell_border(right)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    clear_paragraph(left)
    clear_paragraph(right)

    if PORTRAIT.exists():
        p = left.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        p.add_run().add_picture(str(PORTRAIT), width=Inches(1.18), height=Inches(1.18))

    docx_heading(left, "About me")
    add_docx_text(left, ABOUT, 7.2, "2F333A", False, 6)

    docx_heading(left, "Education")
    for school, major, period in EDUCATION:
        add_docx_text(left, school, 7.2, "2F333A", True, 1)
        add_docx_text(left, major, 7.0, "666B73", False, 1)
        add_docx_text(left, period, 7.0, "666B73", False, 5)

    docx_heading(left, "Skills")
    for label, _level in SKILLS:
        docx_bullet(left, label, 7.1)

    docx_heading(left, "Tools")
    for tool in TOOLS:
        docx_bullet(left, tool, 7.1)

    docx_heading(left, "Language")
    for lang in LANGUAGES:
        docx_bullet(left, lang, 7.1)

    header = right.add_table(rows=1, cols=1)
    set_cell_border(header.cell(0, 0), "26262E", "8")
    set_cell_bg(header.cell(0, 0), "26262E")
    clear_paragraph(header.cell(0, 0))
    add_docx_text(header.cell(0, 0), PROFILE["name"], 19, "FFFFFF", True, 0)
    add_docx_text(header.cell(0, 0), PROFILE["role"], 8, "FFFFFF", True, 0)

    add_docx_text(right, f'{PROFILE["phone"]}    {PROFILE["email"]}', 7.2, "666B73", False, 1)
    add_docx_text(right, f'{PROFILE["website"]}    {PROFILE["location"]}', 7.2, "666B73", False, 5)

    docx_heading(right, "Experience")
    for title, org, period, items in EXPERIENCE:
        add_docx_text(right, f"{title}    {period}", 8.4, "2F333A", True, 1)
        add_docx_text(right, org, 7.4, "666B73", False, 1)
        for item in items:
            docx_bullet(right, item, 7.4)
        add_docx_text(right, "", 1, "FFFFFF", False, 1)

    docx_heading(right, "Practice")
    for title, desc in PRACTICE:
        add_docx_text(right, title, 8.0, "2F333A", True, 1)
        add_docx_text(right, desc, 7.4, "666B73", False, 4)

    docx_heading(right, "Strengths")
    for item in [
        "Quan sát chi tiết, yêu thích sự gọn gàng và nhất quán trong giao diện.",
        "Chủ động học công cụ mới, tiếp thu phản hồi và cải thiện thiết kế qua từng vòng chỉnh sửa.",
        "Có nền tảng kỹ thuật giúp giao tiếp tốt hơn với developer trong quá trình làm sản phẩm.",
    ]:
        docx_bullet(right, item, 7.4)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    draw_pdf()
    draw_docx()
    print(PDF_PATH)
    print(DOCX_PATH)
