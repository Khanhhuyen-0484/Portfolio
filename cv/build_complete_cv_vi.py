from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "cv"
OUT.mkdir(exist_ok=True)

PDF_PATH = OUT / "Khanh_Huyen_CV_Hoan_Chinh.pdf"
DOCX_PATH = OUT / "Khanh_Huyen_CV_Hoan_Chinh.docx"

FONT = Path("C:/Windows/Fonts/arial.ttf")
BOLD = Path("C:/Windows/Fonts/arialbd.ttf")
pdfmetrics.registerFont(TTFont("ArialVN", str(FONT)))
pdfmetrics.registerFont(TTFont("ArialVNBold", str(BOLD)))

NAVY = colors.HexColor("#101828")
BLUE = colors.HexColor("#2563EB")
GRAY = colors.HexColor("#667085")
LIGHT_BLUE = colors.HexColor("#EFF6FF")
LINE = colors.HexColor("#D0D5DD")


def wrap_text(text, font, size, width):
    words = text.split()
    lines, line = [], ""
    for word in words:
        candidate = word if not line else f"{line} {word}"
        if stringWidth(candidate, font, size) <= width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def draw_text(c, x, y, text, size=9, color=NAVY, bold=False, leading=11, width=None):
    font = "ArialVNBold" if bold else "ArialVN"
    c.setFont(font, size)
    c.setFillColor(color)
    lines = wrap_text(text, font, size, width) if width else text.split("\n")
    for line in lines:
        c.drawString(x, y, line)
        y -= leading
    return y


def heading(c, x, y, text, width=1.55 * inch):
    c.setFillColor(BLUE)
    c.setFont("ArialVNBold", 8.7)
    c.drawString(x, y, text.upper())
    c.setStrokeColor(colors.HexColor("#C7D2FE"))
    c.line(x, y - 4, x + width, y - 4)
    return y - 16


def bullet(c, x, y, text, width, size=8.2):
    c.setFillColor(BLUE)
    c.circle(x + 2, y + 2.5, 1.35, fill=1, stroke=0)
    return draw_text(c, x + 10, y, text, size=size, width=width - 10, leading=9.8) - 2


def draw_pdf():
    W, H = letter
    M = 0.42 * inch
    SIDE_W = 2.08 * inch
    GAP = 0.25 * inch
    MAIN_X = M + SIDE_W + GAP
    MAIN_W = W - MAIN_X - M

    c = canvas.Canvas(str(PDF_PATH), pagesize=letter)
    c.setTitle("CV Khanh Huyen")

    c.setFillColor(LIGHT_BLUE)
    c.roundRect(M, M, SIDE_W, H - 2 * M, 14, fill=1, stroke=0)
    c.setFillColor(BLUE)
    c.roundRect(M + 0.18 * inch, H - M - 0.48 * inch, 0.5 * inch, 0.08 * inch, 4, fill=1, stroke=0)

    lx = M + 0.18 * inch
    ly = H - M - 0.78 * inch
    ly = draw_text(c, lx, ly, "KHÁNH\nHUYỀN", 20, bold=True, leading=21)
    ly = draw_text(c, lx, ly - 2, "THỰC TẬP SINH UI/UX DESIGN", 7.55, BLUE, True, 10)
    ly -= 9

    ly = heading(c, lx, ly, "Liên hệ")
    for line in ["TP. Hồ Chí Minh, Việt Nam", "tkhanhhuyen0484@gmail.com"]:
        ly = draw_text(c, lx, ly, line, 7.9, width=SIDE_W - 0.36 * inch, leading=9.5)
    ly -= 8

    ly = heading(c, lx, ly, "Kỹ năng")
    for item in [
        "Thiết kế UI mobile/web",
        "User flow, wireframe",
        "Prototype trên Figma",
        "Visual hierarchy",
        "Design system cơ bản",
        "Usability testing cơ bản",
        "Developer handoff",
        "Responsive design",
    ]:
        ly = bullet(c, lx, ly, item, SIDE_W - 0.34 * inch, 7.75)
    ly -= 8

    ly = heading(c, lx, ly, "Công cụ")
    for item in ["Figma", "Adobe XD", "Photoshop", "Illustrator", "Framer", "Webflow"]:
        ly = bullet(c, lx, ly, item, SIDE_W - 0.34 * inch, 7.75)
    ly -= 8

    ly = heading(c, lx, ly, "Lập trình")
    for item in ["HTML/CSS", "JavaScript", "React", "Tailwind CSS", "Git"]:
        ly = bullet(c, lx, ly, item, SIDE_W - 0.34 * inch, 7.75)
    ly -= 8

    ly = heading(c, lx, ly, "Ngôn ngữ")
    ly = bullet(c, lx, ly, "Tiếng Việt: bản ngữ", SIDE_W - 0.34 * inch, 7.75)
    ly = bullet(c, lx, ly, "Tiếng Anh: đọc hiểu tài liệu chuyên ngành", SIDE_W - 0.34 * inch, 7.75)

    x, y = MAIN_X, H - M - 0.18 * inch
    y = draw_text(c, x, y, "Khánh Huyền", 25, bold=True, leading=27)
    y = draw_text(c, x, y - 2, "UI/UX Designer Intern", 11, BLUE, True, 13, MAIN_W)
    y -= 8

    y = heading(c, x, y, "Mục tiêu nghề nghiệp", MAIN_W)
    y = draw_text(
        c,
        x,
        y,
        "Mong muốn phát triển trong lĩnh vực UI/UX Design, đặc biệt là thiết kế giao diện mobile app và sản phẩm số. Tôi muốn vận dụng nền tảng Công nghệ thông tin, tư duy hệ thống và khả năng thiết kế trực quan để tạo ra các trải nghiệm rõ ràng, dễ sử dụng và có giá trị cho người dùng.",
        8.55,
        width=MAIN_W,
        leading=10.8,
    )
    y -= 8

    y = heading(c, x, y, "Học vấn", MAIN_W)
    y = draw_text(c, x, y, "Đại học Công nghiệp TP. Hồ Chí Minh", 9.8, bold=True, leading=11.2)
    y = draw_text(c, x, y, "Ngành Công nghệ thông tin - Sinh viên năm 4", 8.5, GRAY, leading=10.5)
    y -= 8

    y = heading(c, x, y, "Kinh nghiệm thực hành", MAIN_W)
    sections = [
        (
            "Thiết kế giao diện mobile app",
            [
                "Phân tích nhu cầu người dùng và xây dựng luồng sử dụng cho các chức năng chính.",
                "Thiết kế wireframe, high-fidelity UI và prototype tương tác trên Figma.",
                "Tối ưu bố cục, màu sắc, typography và trạng thái UI để giao diện dễ hiểu hơn.",
            ],
        ),
        (
            "Thiết kế web portfolio cá nhân",
            [
                "Xây dựng cấu trúc website gồm trang Home, Projects, About và Contact.",
                "Tự triển khai giao diện bằng HTML, CSS và JavaScript cơ bản.",
                "Tối ưu responsive layout để nội dung hiển thị tốt trên desktop và mobile.",
            ],
        ),
        (
            "Chuẩn bị handoff và tài liệu thiết kế",
            [
                "Sắp xếp màn hình, component, màu sắc và trạng thái UI có hệ thống.",
                "Diễn giải mục tiêu thiết kế, vấn đề người dùng và giải pháp theo dạng case study.",
            ],
        ),
    ]
    for title, items in sections:
        y = draw_text(c, x, y, title, 9.35, bold=True, leading=11)
        for item in items:
            y = bullet(c, x, y + 1, item, MAIN_W, 8.05)
        y -= 3

    y = heading(c, x, y, "Điểm mạnh", MAIN_W)
    for item in [
        "Có khả năng quan sát chi tiết, yêu thích sự gọn gàng và nhất quán trong giao diện.",
        "Chủ động học công cụ mới, tiếp thu phản hồi và cải thiện thiết kế qua từng vòng chỉnh sửa.",
        "Có nền tảng kỹ thuật giúp giao tiếp tốt hơn với developer trong quá trình handoff.",
    ]:
        y = bullet(c, x, y, item, MAIN_W, 8.05)

    c.setStrokeColor(LINE)
    c.line(M, M + 0.2 * inch, W - M, M + 0.2 * inch)
    c.setFillColor(GRAY)
    c.setFont("ArialVN", 7.5)
    c.drawCentredString(W / 2, M + 0.08 * inch, "Khánh Huyền - CV ứng tuyển UI/UX Design Intern")
    c.save()


def docx_run(run, size=10, color="101828", bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_docx_para(doc, text, size=10, color="101828", bold=False, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    docx_run(r, size, color, bold)
    return p


def add_docx_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    docx_run(r, 11, "2563EB", True)


def add_docx_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run("- " + text)
    docx_run(r, 9.2)


def draw_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)

    add_docx_para(doc, "Khánh Huyền", 24, "101828", True, 0)
    add_docx_para(doc, "UI/UX Designer Intern", 12, "2563EB", True, 6)
    add_docx_para(doc, "TP. Hồ Chí Minh, Việt Nam | tkhanhhuyen0484@gmail.com", 9.5, "667085", False, 8)

    add_docx_heading(doc, "Mục tiêu nghề nghiệp")
    add_docx_para(
        doc,
        "Mong muốn phát triển trong lĩnh vực UI/UX Design, đặc biệt là thiết kế giao diện mobile app và sản phẩm số. Tôi muốn vận dụng nền tảng Công nghệ thông tin, tư duy hệ thống và khả năng thiết kế trực quan để tạo ra các trải nghiệm rõ ràng, dễ sử dụng và có giá trị cho người dùng.",
        9.5,
    )

    add_docx_heading(doc, "Học vấn")
    add_docx_para(doc, "Đại học Công nghiệp TP. Hồ Chí Minh", 10, "101828", True, 1)
    add_docx_para(doc, "Ngành Công nghệ thông tin - Sinh viên năm 4", 9.3, "667085", False, 5)

    add_docx_heading(doc, "Kỹ năng")
    for item in [
        "Thiết kế UI mobile/web, user flow, wireframe và prototype trên Figma.",
        "Tối ưu visual hierarchy, spacing, typography, màu sắc và component consistency.",
        "Design system cơ bản, responsive design, usability testing cơ bản và developer handoff.",
        "Công cụ: Figma, Adobe XD, Photoshop, Illustrator, Framer, Webflow.",
        "Lập trình: HTML/CSS, JavaScript, React, Tailwind CSS, Git.",
    ]:
        add_docx_bullet(doc, item)

    add_docx_heading(doc, "Kinh nghiệm thực hành")
    for title, items in [
        ("Thiết kế giao diện mobile app", [
            "Phân tích nhu cầu người dùng và xây dựng luồng sử dụng cho các chức năng chính.",
            "Thiết kế wireframe, high-fidelity UI và prototype tương tác trên Figma.",
            "Tối ưu bố cục, màu sắc, typography và trạng thái UI để giao diện dễ hiểu hơn.",
        ]),
        ("Thiết kế web portfolio cá nhân", [
            "Xây dựng cấu trúc website gồm trang Home, Projects, About và Contact.",
            "Tự triển khai giao diện bằng HTML, CSS và JavaScript cơ bản.",
            "Tối ưu responsive layout để nội dung hiển thị tốt trên desktop và mobile.",
        ]),
        ("Chuẩn bị handoff và tài liệu thiết kế", [
            "Sắp xếp màn hình, component, màu sắc và trạng thái UI có hệ thống.",
            "Diễn giải mục tiêu thiết kế, vấn đề người dùng và giải pháp theo dạng case study.",
        ]),
    ]:
        add_docx_para(doc, title, 10, "101828", True, 1)
        for item in items:
            add_docx_bullet(doc, item)

    add_docx_heading(doc, "Điểm mạnh")
    for item in [
        "Có khả năng quan sát chi tiết, yêu thích sự gọn gàng và nhất quán trong giao diện.",
        "Chủ động học công cụ mới, tiếp thu phản hồi và cải thiện thiết kế qua từng vòng chỉnh sửa.",
        "Có nền tảng kỹ thuật giúp giao tiếp tốt hơn với developer trong quá trình handoff.",
    ]:
        add_docx_bullet(doc, item)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    draw_pdf()
    draw_docx()
    print(PDF_PATH)
    print(DOCX_PATH)
